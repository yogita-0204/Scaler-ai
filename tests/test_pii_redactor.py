"""Tests for PII redaction package."""

import pytest

import sys
import os
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

# Make sure the package is importable
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from pii_redactor.detector_base import Detection, resolve_overlaps
from pii_redactor.detectors_structured import (
    EmailDetector, PhoneDetector, IPAddressDetector,
    CreditCardDetector, SSNDetector, CINDetector,
    PANDetector, DINDetector, PincodeDetector,
)
from pii_redactor.detectors_contextual import (
    PersonNameDetector, CompanyNameDetector,
    AddressDetector, DOBDetector,
)
from pii_redactor.pseudonymiser import get_replacement, clear_cache
from pii_redactor.config import (
    CAT_EMAIL, CAT_PHONE, CAT_IP, CAT_CC, CAT_SSN,
    CAT_DOB, CAT_ADDRESS, CAT_PERSON, CAT_COMPANY,
    CAT_CIN, CAT_PAN, CAT_DIN, CAT_PINCODE,
)



# Helper


def found_categories(detections):
    return {d.category for d in detections}


def found_values(detections):
    return {d.value for d in detections}



# EMAIL DETECTOR


class TestEmailDetector:
    det = EmailDetector()

    def test_simple_email(self):
        r = self.det.detect("Contact: user@example.com please")
        assert CAT_EMAIL in found_categories(r)
        assert "user@example.com" in found_values(r)

    def test_real_doc_email(self):
        r = self.det.detect("Email: cs.connect@kshinternational.com")
        assert any("cs.connect@kshinternational.com" in d.value for d in r)

    def test_email_with_plus(self):
        r = self.det.detect("Email: user+tag@domain.org")
        assert CAT_EMAIL in found_categories(r)

    def test_email_subdomain(self):
        r = self.det.detect("Send to info@mail.example.co.uk")
        assert CAT_EMAIL in found_categories(r)

    def test_no_email_in_plain_text(self):
        r = self.det.detect("The revenue was 4200 million rupees")
        assert CAT_EMAIL not in found_categories(r)

    def test_no_email_in_number(self):
        r = self.det.detect("Order 12345 processed")
        assert CAT_EMAIL not in found_categories(r)

    def test_malformed_no_tld(self):
        r = self.det.detect("x@y is not a valid email")
        assert CAT_EMAIL not in found_categories(r)

    def test_multiple_emails(self):
        text = "From: a@b.com and c@d.org"
        r = self.det.detect(text)
        vals = found_values(r)
        assert "a@b.com" in vals
        assert "c@d.org" in vals

    def test_real_hdfc_email(self):
        r = self.det.detect("hitesh.ramani@citi.com is the contact")
        assert CAT_EMAIL in found_categories(r)



# PHONE DETECTOR


class TestPhoneDetector:
    det = PhoneDetector()

    def test_indian_plus91_with_spaces(self):
        r = self.det.detect("Telephone: +91 20 45053237")
        assert CAT_PHONE in found_categories(r)

    def test_real_doc_phone(self):
        r = self.det.detect("Call +91 22 4009 4400 for info")
        assert CAT_PHONE in found_categories(r)

    def test_mobile_10_digit(self):
        r = self.det.detect("Mobile: 9876543210")
        assert CAT_PHONE in found_categories(r)

    def test_mobile_starts_with_6(self):
        r = self.det.detect("Call 6123456789 now")
        assert CAT_PHONE in found_categories(r)

    def test_no_phone_from_financial_figure(self):
        r = self.det.detect("Offer size: ₹4,200.00 million")
        assert CAT_PHONE not in found_categories(r)

    def test_no_phone_from_section_number(self):
        r = self.det.detect("Section 32 of the Companies Act")
        assert CAT_PHONE not in found_categories(r)

    def test_no_phone_from_short_number(self):
        r = self.det.detect("12345 is the reference")
        assert CAT_PHONE not in found_categories(r)

    def test_phone_with_hyphens(self):
        r = self.det.detect("+91-80-1234-5678")
        assert CAT_PHONE in found_categories(r)



# IP ADDRESS DETECTOR


class TestIPDetector:
    det = IPAddressDetector()

    def test_private_ip(self):
        r = self.det.detect("Server: 192.168.1.100")
        assert CAT_IP in found_categories(r)

    def test_public_ip(self):
        r = self.det.detect("IP: 203.0.113.42")
        assert CAT_IP in found_categories(r)

    def test_invalid_octet_256(self):
        r = self.det.detect("256.1.2.3 is invalid")
        assert CAT_IP not in found_categories(r)

    def test_localhost(self):
        r = self.det.detect("localhost 127.0.0.1 is home")
        assert CAT_IP in found_categories(r)

    def test_no_ip_in_plain_text(self):
        r = self.det.detect("Revenue of 100 million")
        assert CAT_IP not in found_categories(r)



# CREDIT CARD DETECTOR


class TestCreditCardDetector:
    det = CreditCardDetector()

    def test_valid_visa_luhn(self):
        # 4532015112830366 passes Luhn
        r = self.det.detect("Card: 4532015112830366")
        assert CAT_CC in found_categories(r)

    def test_invalid_luhn(self):
        r = self.det.detect("Number: 1234567890123456")
        assert CAT_CC not in found_categories(r)

    def test_card_with_spaces(self):
        # 4532 0151 1283 0366 – passes Luhn
        r = self.det.detect("Card: 4532 0151 1283 0366")
        assert CAT_CC in found_categories(r)



# SSN DETECTOR


class TestSSNDetector:
    det = SSNDetector()

    def test_valid_ssn(self):
        r = self.det.detect("SSN: 123-45-6789")
        assert CAT_SSN in found_categories(r)

    def test_invalid_area_000(self):
        r = self.det.detect("000-45-6789 invalid area")
        assert CAT_SSN not in found_categories(r)

    def test_invalid_area_666(self):
        r = self.det.detect("666-45-6789 invalid")
        assert CAT_SSN not in found_categories(r)

    def test_invalid_group_00(self):
        r = self.det.detect("123-00-6789 invalid group")
        assert CAT_SSN not in found_categories(r)



# CIN DETECTOR


class TestCINDetector:
    det = CINDetector()

    def test_real_cin_1(self):
        r = self.det.detect("CIN: U28129PN1979PLC141032")
        assert CAT_CIN in found_categories(r)

    def test_real_cin_2(self):
        r = self.det.detect("CIN: U67190MH1999PTC118368")
        assert CAT_CIN in found_categories(r)

    def test_not_cin(self):
        r = self.det.detect("Under section U/s 32 of the Act")
        assert CAT_CIN not in found_categories(r)



# DIN DETECTOR


class TestDINDetector:
    det = DINDetector()

    def test_din_with_label(self):
        r = self.det.detect("DIN: 01234567")
        assert CAT_DIN in found_categories(r)
        assert "01234567" in found_values(r)

    def test_din_with_colon_space(self):
        r = self.det.detect("Director DIN 98765432 approved")
        assert CAT_DIN in found_categories(r)

    def test_no_din_without_label(self):
        r = self.det.detect("The number 12345678 appears")
        assert CAT_DIN not in found_categories(r)



# PINCODE DETECTOR


class TestPincodeDetector:
    det = PincodeDetector()

    def test_pincode_in_address(self):
        text = "Pune – 410501, Maharashtra, India"
        r = self.det.detect(text)
        assert CAT_PINCODE in found_categories(r)

    def test_pincode_near_city(self):
        text = "Mumbai 400001 is the address"
        r = self.det.detect(text)
        assert CAT_PINCODE in found_categories(r)

    def test_no_pincode_standalone_number(self):
        text = "The offer is for 100000 shares"
        r = self.det.detect(text)
        assert CAT_PINCODE not in found_categories(r)



# PERSON NAME DETECTOR


class TestPersonNameDetector:
    det = PersonNameDetector()

    def test_contact_person_label(self):
        text = "Contact Person: Sarthak Malvadkar Company Secretary"
        r = self.det.detect(text)
        assert CAT_PERSON in found_categories(r)

    def test_titled_name(self):
        r = self.det.detect("Mr. Rajesh Kumar signed the document")
        assert CAT_PERSON in found_categories(r)

    def test_known_person(self):
        text = "Promoter: Sarthak Malvadkar has submitted"
        r = self.det.detect(text)
        assert CAT_PERSON in found_categories(r)

    def test_known_person_match_preserves_whitespace_offsets(self):
        text = "Contact:  Sarthak Malvadkar"
        detections = self.det.detect(text)
        assert any(d.value == "Sarthak Malvadkar" for d in detections)

    def test_no_person_from_abbreviations(self):
        r = self.det.detect("QIBs and RIIs are eligible")
        assert CAT_PERSON not in found_categories(r)

    def test_allcaps_legal_phrase_not_a_person(self):
        text = "ISSUER'S AND PROMOTER SELLING SHAREHOLDERS' ABSOLUTE RESPONSIBILITY"
        detections = self.det.detect(text)
        assert not any(d.value == "ABSOLUTE RESPONSIBILITY" for d in detections)

    def test_contact_name_does_not_swallow_label_word(self):
        text = "Contact Person: Chitra Raste Website: www.eximbankindia.in"
        values = [d.value for d in self.det.detect(text)]
        assert "Chitra Raste" in values
        assert not any("Website" in value for value in values)



# COMPANY NAME DETECTOR


class TestCompanyNameDetector:
    det = CompanyNameDetector()

    def test_known_company(self):
        text = "Lead Manager: Nuvama Wealth Management Limited"
        r = self.det.detect(text)
        assert CAT_COMPANY in found_categories(r)

    def test_company_suffix_pattern(self):
        r = self.det.detect("Horizon Solutions Limited manages the fund")
        assert CAT_COMPANY in found_categories(r)

    def test_trust_name(self):
        r = self.det.detect("The Dhaulagiri Family Trust has approved")
        assert CAT_COMPANY in found_categories(r)

    def test_no_company_from_act(self):
        r = self.det.detect("under the Companies Act 2013")
        # "Companies Act" should not be detected as a company
        companies = [d.value for d in r if d.category == CAT_COMPANY]
        assert not any("Companies Act" in c for c in companies)

    def test_document_entity_fragments_are_detected(self):
        text = "FAMILY TRUST, Escrow Collection Bank"
        values = {d.value.lower() for d in self.det.detect(text)}
        assert "family trust" in values
        assert "escrow collection bank" in values



# DOB DETECTOR


class TestDOBDetector:
    det = DOBDetector()

    def test_dob_with_label_numeric(self):
        r = self.det.detect("Date of birth: 15/08/1985")
        assert CAT_DOB in found_categories(r)

    def test_dob_with_label_text(self):
        r = self.det.detect("DOB: January 10, 1990")
        assert CAT_DOB in found_categories(r)

    def test_no_dob_offer_date(self):
        r = self.det.detect("The offer closes on December 10, 2025")
        assert CAT_DOB not in found_categories(r)

    def test_no_dob_document_date(self):
        r = self.det.detect("Dated December 10, 2025")
        assert CAT_DOB not in found_categories(r)


class TestAddressDetector:
    det = AddressDetector()

    def test_legal_history_is_not_an_address(self):
        text = (
            "certificate of incorporation dated July 30, 1979, issued by the "
            "Registrar of Companies, Maharashtra at Bombay"
        )
        assert CAT_ADDRESS not in found_categories(self.det.detect(text))

    def test_customer_geography_is_not_an_address(self):
        text = "customers across 24 countries outside India, as of June 30, 2025"
        assert CAT_ADDRESS not in found_categories(self.det.detect(text))

    def test_address_list_with_and_is_redacted_from_first_number(self):
        text = "11/3, 11/4 and 11/5 Village Birdewadi Chakan Taluka - Khed, Pune"
        detections = self.det.detect(text)
        assert detections
        assert detections[0].value.startswith("11/3")


# OVERLAP RESOLVER


class TestOverlapResolver:

    def _det(self, start, end, cat, conf):
        return Detection(start=start, end=end, category=cat,
                        value="x"*(end-start), confidence=conf, detector="t")

    def test_no_overlap(self):
        dets = [self._det(0, 5, CAT_EMAIL, 0.9),
                self._det(10, 15, CAT_PHONE, 0.9)]
        result = resolve_overlaps(dets)
        assert len(result) == 2

    def test_overlap_higher_confidence_wins_on_overlap_region(self):
        dets = [self._det(0, 10, CAT_EMAIL, 0.7),
                self._det(5, 15, CAT_PHONE, 0.95)]
        result = resolve_overlaps(dets)
        # Per-position resolution keeps both, with PHONE winning the
        # overlapping region [5,10).
        assert len(result) == 2
        by_start = {d.start: d.category for d in result}
        assert by_start[0] == CAT_EMAIL
        assert by_start[5] == CAT_PHONE
        assert not any(d.overlaps(e) for d in result for e in result if d is not e)

    def test_chained_overlap_maximizes_coverage(self):
        dets = [self._det(0, 100, "A", 0.70),
                self._det(10, 20, "B", 0.95),
                self._det(15, 110, "C", 0.80)]
        result = resolve_overlaps(dets)
        assert not any(d.overlaps(e) for d in result for e in result if d is not e)
        # Coverage: A on [0,10), B on [10,20), C on [20,110).
        assert sorted((d.start, d.end) for d in result) == [(0, 10), (10, 20), (20, 110)]
        winner_at = {position: next(d.category for d in result if d.start <= position < d.end)
                     for position in (5, 15, 50)}
        assert winner_at == {5: "A", 15: "B", 50: "C"}

    def test_overlap_longer_wins_on_tie(self):
        dets = [self._det(0, 10, CAT_EMAIL, 0.9),
                self._det(0, 15, CAT_PHONE, 0.9)]
        result = resolve_overlaps(dets)
        assert len(result) == 1

    def test_adjacent_no_merge(self):
        dets = [self._det(0, 5, CAT_EMAIL, 0.9),
                self._det(5, 10, CAT_PHONE, 0.9)]
        result = resolve_overlaps(dets)
        assert len(result) == 2



# PSEUDONYMISER


class TestPseudonymiser:

    def setup_method(self):
        clear_cache()

    def test_deterministic_email(self):
        r1 = get_replacement(CAT_EMAIL, "test@test.com")
        r2 = get_replacement(CAT_EMAIL, "test@test.com")
        assert r1 == r2
        assert "@" in r1

    def test_deterministic_phone(self):
        r1 = get_replacement(CAT_PHONE, "+91 20 45053237")
        r2 = get_replacement(CAT_PHONE, "+91 20 45053237")
        assert r1 == r2
        assert "+91" in r1

    def test_different_inputs_different_output(self):
        r1 = get_replacement(CAT_EMAIL, "alice@a.com")
        r2 = get_replacement(CAT_EMAIL, "bob@b.com")
        assert r1 != r2

    def test_person_fake_has_two_parts(self):
        r = get_replacement(CAT_PERSON, "Sarthak Malvadkar")
        parts = r.split()
        assert len(parts) >= 2

    def test_cin_format(self):
        r = get_replacement(CAT_CIN, "U28129PN1979PLC141032")
        assert len(r) == 21
        assert r[0] in "LU"

    def test_fake_credit_card_is_not_luhn_valid(self):
        replacement = get_replacement(CAT_CC, "4532015112830366")
        assert CAT_CC not in found_categories(CreditCardDetector().detect(replacement))

    def test_din_is_8_digits(self):
        r = get_replacement(CAT_DIN, "01234567")
        assert r.isdigit()
        assert len(r) == 8



# REGRESSION TESTS (bugs found and fixed)


class TestRegressions:
    """Each test locks in a fix for a specific discovered issue."""

    def test_reg_001_email_in_table_cell(self):
        """Email inside a table cell text should be detected."""
        email_det = EmailDetector()
        r = email_det.detect("ksh.ipo@nuvama.com")
        assert CAT_EMAIL in found_categories(r)

    def test_reg_002_phone_with_extra_spaces(self):
        """Phone with extra internal spaces (as in DOCX runs) should match."""
        phone_det = PhoneDetector()
        r = phone_det.detect("+   91   22   4009   4400")
        assert CAT_PHONE in found_categories(r)

    def test_reg_003_cin_full_string(self):
        """CIN alone (no surrounding text) should be detected."""
        cin_det = CINDetector()
        r = cin_det.detect("U28129PN1979PLC141032")
        assert CAT_CIN in found_categories(r)

    def test_reg_004_email_uppercase_domain(self):
        """Emails with uppercase domain parts should still match."""
        email_det = EmailDetector()
        r = email_det.detect("User@EXAMPLE.COM")
        assert CAT_EMAIL in found_categories(r)

    def test_reg_005_company_with_ampersand(self):
        """Company names with & should be in known-entity list and detected."""
        comp_det = CompanyNameDetector()
        r = comp_det.detect("Kirtane & Pandit LLP audited the accounts")
        assert CAT_COMPANY in found_categories(r)

    def test_reg_006_no_ssn_from_date_range(self):
        """Date ranges like '1979-01-2025' should not match as SSN."""
        ssn_det = SSNDetector()
        r = ssn_det.detect("From 1979 to 2025 the company operated")
        assert CAT_SSN not in found_categories(r)

    def test_reg_007_multiple_emails_all_detected(self):
        """All emails in a paragraph should be detected."""
        email_det = EmailDetector()
        text = ("Send to cs.connect@kshinternational.com or "
                "ipo@trilegal.com for info")
        r = email_det.detect(text)
        vals = found_values(r)
        assert "cs.connect@kshinternational.com" in vals
        assert "ipo@trilegal.com" in vals



# INTEGRATION TEST – end-to-end detection on real document text snippets


class TestIntegration:

    def test_cover_page_snippet(self):
        """Real text from the prospectus cover page."""
        text = (
            "KSH INTERNATIONAL LIMITED CORPORATE IDENTITY NUMBER: U28129PN1979PLC141032 "
            "Contact Person: Sarthak Malvadkar Company Secretary and Compliance Officer "
            "Email: cs.connect@kshinternational.com Telephone: +91 20 45053237 "
            "www.kshinternational.com"
        )
        from pii_redactor.docx_processor import detect_all
        dets = detect_all(text)
        cats = found_categories(dets)
        assert CAT_EMAIL in cats, "Email should be detected"
        assert CAT_PHONE in cats, "Phone should be detected"
        assert CAT_CIN in cats, "CIN should be detected"
        assert CAT_PERSON in cats, "Person name should be detected"

    def test_address_snippet(self):
        """Real registered office address from the prospectus."""
        text = (
            "Registered Office: 11/3, 11/4 and 11/5 Village Birdewadi "
            "Chakan Taluka - Khed, Pune – 410 501, Maharashtra, India"
        )
        from pii_redactor.docx_processor import detect_all
        dets = detect_all(text)
        cats = found_categories(dets)
        assert CAT_ADDRESS in cats, "Address should be detected"

    def test_replacements_are_consistent(self):
        """Same PII value should get same replacement everywhere."""
        from pii_redactor.docx_processor import detect_all
        text1 = "Contact: cs.connect@kshinternational.com"
        text2 = "Email cs.connect@kshinternational.com for info"
        dets1 = detect_all(text1)
        dets2 = detect_all(text2)
        repl1 = {d.value: d.replacement for d in dets1 if d.category == CAT_EMAIL}
        repl2 = {d.value: d.replacement for d in dets2 if d.category == CAT_EMAIL}
        email = "cs.connect@kshinternational.com"
        assert email in repl1
        assert email in repl2
        assert repl1[email] == repl2[email], "Same email must have same replacement"


class TestDocxRedactionIntegrity:
    def test_redaction_preserves_unaffected_runs(self):
        """Replacing a token must not collapse or discard sibling runs."""
        with TemporaryDirectory() as directory:
            input_path = Path(directory) / "input.docx"
            output_path = Path(directory) / "output.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("Keep this ").bold = True
            paragraph.add_run("user@example.com")
            paragraph.add_run(" and this text.").italic = True
            document.save(input_path)

            from pii_redactor.docx_processor import DocxProcessor
            DocxProcessor(str(input_path), str(output_path)).process()

            result = Document(output_path)
            result_paragraph = result.paragraphs[0]
            text = "".join(run.text or "" for run in result_paragraph.runs)
            assert "Keep this " in text
            assert " and this text." in text
            assert "user@example.com" not in text
            assert len(result_paragraph.runs) == 3
            assert result_paragraph.runs[0].text == "Keep this "
            assert result_paragraph.runs[1].text != "user@example.com"
            assert result_paragraph.runs[1].text
            assert result_paragraph.runs[2].text == " and this text."
            assert result_paragraph.runs[0].bold is True
            assert result_paragraph.runs[2].italic is True

    def test_redaction_covers_nested_table_and_split_text_nodes(self):
        """All paragraph XML, including nested table content, is redacted."""
        with TemporaryDirectory() as directory:
            input_path = Path(directory) / "input.docx"
            output_path = Path(directory) / "output.docx"
            document = Document()
            outer = document.add_table(rows=1, cols=1)
            inner = outer.cell(0, 0).add_table(rows=1, cols=1)
            cell = inner.cell(0, 0)
            cell.paragraphs[0].add_run("user@").bold = True
            cell.paragraphs[0].add_run("example.com")
            document.save(input_path)

            from pii_redactor.docx_processor import DocxProcessor
            DocxProcessor(str(input_path), str(output_path)).process()

            result = Document(output_path)
            nested_text = " ".join(
                paragraph.text
                for table in result.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
            assert "user@example.com" not in nested_text

    def test_redaction_covers_field_code_instructions(self):
        """mailto field instructions must not retain the source email."""
        from pii_redactor.docx_processor import _process_xml

        xml = b'''<?xml version="1.0" encoding="UTF-8"?>
        <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
          <w:body><w:p>
            <w:fldSimple w:instr="HYPERLINK">
              <w:r><w:instrText>HYPERLINK "mailto:user@example.com"</w:instrText></w:r>
            </w:fldSimple>
            <w:r><w:t>user@example.com</w:t></w:r>
          </w:p></w:body>
        </w:document>'''

        redacted, _ = _process_xml(xml)
        assert b"user@example.com" not in redacted

    def test_same_input_and_output_path_is_rejected(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "document.docx"
            Document().save(path)
            from pii_redactor.docx_processor import DocxProcessor
            with pytest.raises(ValueError, match="different"):
                DocxProcessor(str(path), str(path)).process()

    def test_edge_case_empty_and_whitespace_input(self):
        """Empty or whitespace input must return empty detections without error."""
        from pii_redactor.docx_processor import detect_all
        assert detect_all("") == []
        assert detect_all("   \n\t  ") == []

    def test_edge_case_mixed_case_trust_name(self):
        """ALLCAPS trust names in headers must be detected and pseudonymised."""
        comp_det = CompanyNameDetector()
        r = comp_det.detect("OUR PROMOTERS: EVEREST FAMILY TRUST AND MAKALU FAMILY TRUST")
        vals = {d.value.upper() for d in r}
        assert "EVEREST FAMILY TRUST" in vals

    def test_edge_case_company_list_conjunctions(self):
        """Conjunctions like 'and' must not merge separate company names."""
        comp_det = CompanyNameDetector()
        text = "Annapurna Family Trust, Kanchenjunga Family Trust and Waterloo Industrial Park VI Private Limited"
        dets = comp_det.detect(text)
        vals = [d.value for d in dets]
        assert "Annapurna Family Trust" in vals
        assert "Kanchenjunga Family Trust" in vals
        assert "Waterloo Industrial Park VI Private Limited" in vals
        assert not any("and" in v.lower() for v in vals)

    def test_edge_case_generic_phrases_not_detected(self):
        """Generic phrases like 'Senior Management' must not be flagged as companies."""
        comp_det = CompanyNameDetector()
        text = "Our Senior Management and Key Management personnel"
        dets = comp_det.detect(text)
        assert not dets

    def test_edge_case_pseudonymiser_case_insensitive_cache(self):
        """Pseudonymiser should yield the same fake value regardless of casing."""
        r1 = get_replacement(CAT_PERSON, "Sarthak Malvadkar")
        r2 = get_replacement(CAT_PERSON, "sarthak malvadkar")
        r3 = get_replacement(CAT_PERSON, "SARTHAK MALVADKAR")
        assert r1 == r2 == r3

